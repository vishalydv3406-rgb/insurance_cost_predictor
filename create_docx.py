import docx
import re
from docx.shared import Pt, Inches, RGBColor

def apply_formatting(paragraph, text):
    # This is a very basic markdown formatting handler
    # We will split text by ** for bold
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # check for italic *
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sub_part in sub_parts:
                if sub_part.startswith('*') and sub_part.endswith('*') and len(sub_part) > 2:
                    run = paragraph.add_run(sub_part[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sub_part)

def md_to_docx(md_filepath, docx_filepath):
    doc = docx.Document()
    
    with open(md_filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('---'):
            continue
        
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            apply_formatting(p, line[2:])
        elif line.startswith('| '): # Simple table-ish output
            p = doc.add_paragraph()
            apply_formatting(p, line)
            p.paragraph_format.left_indent = Inches(0.5)
        elif line.startswith('👉 '):
            p = doc.add_paragraph()
            run = p.add_run('👉 ')
            apply_formatting(p, line[2:])
        else:
            p = doc.add_paragraph()
            apply_formatting(p, line)
            
    doc.save(docx_filepath)

if __name__ == "__main__":
    md_to_docx('Project_Report.md', 'Insurance_Project_Report.docx')
    print("Docx created successfully.")
